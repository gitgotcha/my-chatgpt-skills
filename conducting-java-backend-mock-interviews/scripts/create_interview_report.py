"""Build a printable Java backend mock-interview review report from JSON."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TABLE_FILL = "E8EEF5"
TEXT_COLOR = RGBColor(0, 0, 0)
TABLE_WIDTH_DXA = 9360


def _set_run_font(run: Any, *, size: float | None = None, bold: bool | None = None,
                  color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_width(cell: Any, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    table_properties.append(table_indent)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            _set_cell_width(cell, width)


def _write_cell(cell: Any, value: object, *, header: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(value if value not in (None, "") else "未提供"))
    _set_run_font(run, size=10.5, bold=header, color=TEXT_COLOR)
    if header:
        _set_cell_shading(cell, TABLE_FILL)


def _add_page_number(paragraph: Any) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _add_toc_field(paragraph: Any) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(field)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("Java 后端模拟面试报告 | 第 ")
    _set_run_font(footer_run, size=9, color=RGBColor(89, 89, 89))
    _add_page_number(footer)
    footer_run = footer.add_run(" 页")
    _set_run_font(footer_run, size=9, color=RGBColor(89, 89, 89))

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_key_value_table(document: Document, items: dict[str, object]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in items.items():
        cells = table.add_row().cells
        _write_cell(cells[0], key, header=True)
        _write_cell(cells[1], value)
    _set_table_geometry(table, [2700, TABLE_WIDTH_DXA - 2700])


def add_score_table(document: Document, scores: list[dict[str, object]]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("维度", "得分", "评价"), strict=True):
        _write_cell(cell, label, header=True)
    for score in scores:
        cells = table.add_row().cells
        _write_cell(cells[0], score.get("dimension", "未命名维度"))
        _write_cell(cells[1], score.get("score", "未评分"))
        _write_cell(cells[2], score.get("comment", "未提供"))
    _set_table_geometry(table, [2200, 1100, 6060])


def _add_value(document: Document, value: object) -> None:
    if value in (None, "", [], {}):
        document.add_paragraph("本场未覆盖。")
    elif isinstance(value, str):
        document.add_paragraph(value)
    elif isinstance(value, list):
        for item in value:
            document.add_paragraph(str(item), style="List Bullet")
    elif isinstance(value, dict):
        add_key_value_table(document, value)
    else:
        document.add_paragraph(str(value))


def _add_questions(document: Document, questions: list[dict[str, object]]) -> None:
    for index, item in enumerate(questions, start=1):
        identifier = item.get("id", index)
        add_heading(document, f"{identifier}. 面试官问题：{item.get('question', '未记录')}", level=2)
        fields = (
            ("用户回答摘要", "answer_summary"),
            ("回答评价", "evaluation"),
            ("标准答案", "standard_answer"),
            ("推荐口述版本", "spoken_answer"),
            ("表述分析", "expression_feedback"),
        )
        for title, key in fields:
            paragraph = document.add_paragraph()
            label = paragraph.add_run(f"{title}：")
            _set_run_font(label, bold=True, color=RGBColor.from_string(DARK_BLUE))
            paragraph.add_run(str(item.get(key, "未提供")))
        document.add_paragraph("可能的下一层追问：")
        follow_ups = item.get("next_questions", [])
        _add_value(document, follow_ups if isinstance(follow_ups, list) else [follow_ups])


def create_report(input_path: Path, output_path: Path) -> None:
    data = json.loads(input_path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("interview record JSON root must be an object")
    if data.get("state") == "review_pending":
        raise ValueError("pending mock handoff must be reviewed before a report is generated")

    document = Document()
    _configure_document(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run("Java 后端模拟面试报告")
    _set_run_font(title_run, size=24, bold=True, color=RGBColor.from_string(DARK_BLUE))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(54)
    subtitle_run = subtitle.add_run("中大型互联网公司岗位训练与复盘")
    _set_run_font(subtitle_run, size=12, color=RGBColor(89, 89, 89))
    document.add_page_break()

    add_heading(document, "目录", level=1)
    toc = document.add_paragraph("在 Word 中打开后将自动更新目录。")
    _add_toc_field(toc)
    document.add_page_break()

    basic_info = data.get("basic_info", {})
    if not isinstance(basic_info, dict):
        basic_info = {"原始信息": basic_info}
    if basic_info.get("completed") is False:
        basic_info = {**basic_info, "面试完成度说明": "本次面试未完整进行，以下评分基于已完成内容。"}

    add_heading(document, "一、面试基本信息")
    add_key_value_table(document, basic_info)
    add_heading(document, "二、面试结论")
    _add_value(document, data.get("conclusion"))
    add_heading(document, "三、分项评分")
    scores = data.get("scores", [])
    add_score_table(document, scores if isinstance(scores, list) else [])

    for heading, key in (
        ("四、简历风险扫描", "risk_scan"),
        ("五、项目掌握度分析", "project_analysis"),
        ("六、算法表现", "algorithm_performance"),
        ("七、主要优势", "strengths"),
        ("八、主要问题", "issues"),
    ):
        add_heading(document, heading)
        _add_value(document, data.get(key))

    add_heading(document, "九、逐题复盘")
    questions = data.get("questions", [])
    _add_questions(document, questions if isinstance(questions, list) else [])
    add_heading(document, "十、薄弱知识点清单")
    _add_value(document, data.get("knowledge_gaps"))
    add_heading(document, "十一、下一次训练建议")
    _add_value(document, data.get("training_plan"))
    add_heading(document, "附录：推荐答案与代码")
    _add_value(document, data.get("appendix"))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="UTF-8 interview-record JSON file")
    parser.add_argument("output", type=Path, help="output .docx path")
    arguments = parser.parse_args()
    create_report(arguments.input, arguments.output)


if __name__ == "__main__":
    main()
