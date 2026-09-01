from __future__ import annotations

import json
import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.create_review_report import create_review_report


USER_ID = "11111111-1111-4111-8111-111111111111"
SESSION_EVENT_ID = "22222222-2222-4222-8222-222222222222"
REVIEW_EVENT_ID = "33333333-3333-4333-8333-333333333333"
SESSION_ID = "REAL-20260814T000000Z-aaaaaaaa-aaaa-4aaa-8aaa-aaaaaaaaaaaa"


def _review_event() -> dict[str, object]:
    """One persisted schema-1.2 ``interview.review.completed`` JSON record."""
    return {
        "schemaVersion": "1.2",
        "eventId": REVIEW_EVENT_ID,
        "eventKey": f"{USER_ID}:interview:review:{SESSION_ID}:v1",
        "eventType": "interview.review.completed",
        "userId": USER_ID,
        "username": "乔炳源",
        "sessionId": SESSION_ID,
        "interviewType": "real",
        "domain": "java-backend",
        "completedAt": "2026-08-14T01:00:00Z",
        "reviewVersion": 1,
        "sourceSessionEventId": SESSION_EVENT_ID,
        "sourceType": "real",
        "evidenceType": "full_transcript",
        "evidenceConfidence": "high",
        "questionReviews": [{
            "questionId": "Q-001",
            "assessment": "缓存一致性方案只答到删除缓存，缺少双删与重试。",
            "evidence": {"source": "answer"},
            "recommendations": ["补强缓存一致性"],
        }],
        "profileChanges": [{
            "kind": "weakness",
            "outcome": "failed",
            "domain": "java-backend",
            "weaknessId": "W-001",
            "evidenceRefs": ["Q-001"],
        }],
        "recommendations": ["复测缓存一致性"],
        "applyProfileChanges": False,
        "persistenceStatus": "cloud_accepted",
        "overallAssessment": "整体可用，缓存一致性需要补强。",
    }


class ReviewReportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.directory = Path(tempfile.mkdtemp())
        self.addCleanup(shutil.rmtree, self.directory, True)
        self.source = self.directory / "review.json"
        self.output = self.directory / "review-report.docx"

    def _render(self, payload: dict[str, object]) -> str:
        self.source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

        create_review_report(self.source, self.output)

        document = Document(self.output)
        return "\n".join(paragraph.text for paragraph in document.paragraphs) + "\n" + "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )

    def test_report_identity_uses_username_and_user_id(self) -> None:
        text = self._render(_review_event())

        self.assertTrue(self.output.exists())
        self.assertIn("乔炳源", text)
        self.assertIn(USER_ID, text)
        self.assertIn(SESSION_ID, text)
        self.assertIn(SESSION_EVENT_ID, text)

    def test_report_renders_evidence_assessment_changes_and_recommendations(self) -> None:
        text = self._render(_review_event())

        self.assertIn("证据类型", text)
        self.assertIn("full_transcript", text)
        self.assertIn("缓存一致性方案只答到删除缓存", text)
        self.assertIn("复测缓存一致性", text)
        self.assertIn("weakness", text)

    def test_report_keeps_navigation_and_pagination(self) -> None:
        self._render(_review_event())

        document = Document(self.output)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("目录", text)
        self.assertIn("PAGE", document.sections[0].footer._element.xml)

    def test_legacy_report_json_is_rejected_instead_of_rendered(self) -> None:
        legacy = {
            "basic_info": {"candidate": "候选人", "completed": False},
            "record_quality": {"overall_confidence": 0.8},
            "overall_assessment": {"summary": "需要补强缓存一致性。"},
        }
        self.source.write_text(json.dumps(legacy, ensure_ascii=False), encoding="utf-8")

        with self.assertRaisesRegex(ValueError, "schemaVersion 1.2"):
            create_review_report(self.source, self.output)

        self.assertFalse(self.output.exists())

    def test_non_object_root_is_rejected(self) -> None:
        self.source.write_text("[]", encoding="utf-8")

        with self.assertRaisesRegex(ValueError, "must be an object"):
            create_review_report(self.source, self.output)


if __name__ == "__main__":
    unittest.main()
