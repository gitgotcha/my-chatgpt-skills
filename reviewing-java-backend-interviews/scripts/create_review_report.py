"""Create a local Word report from one persisted schema-1.2 review JSON."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _font(run: Any, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    _append_field(footer, "PAGE")
    footer.add_run(" 页")


def _append_field(paragraph: Any, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instruction_text, separate, end))


def _add_toc(document: Document) -> None:
    document.add_heading("目录", level=1)
    paragraph = document.add_paragraph()
    _append_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    document.add_page_break()


def _value(document: Document, value: object) -> None:
    if value in (None, "", [], {}):
        document.add_paragraph("本场未覆盖。")
    elif isinstance(value, str):
        document.add_paragraph(value)
    elif isinstance(value, list):
        for item in value:
            document.add_paragraph(str(item), style="List Bullet")
    elif isinstance(value, dict):
        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for key, item in value.items():
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(item)
    else:
        document.add_paragraph(str(value))


def _schema12_report(document: Document, data: dict[str, object]) -> None:
    """Render only the persisted review JSON, without consulting a snapshot."""
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("Java 后端面试复盘报告")
    _font(run, size=24, bold=True)
    document.add_page_break()
    _add_toc(document)

    document.add_heading("一、身份与会话信息", level=1)
    identity = {
        "姓名": data.get("username", "未提供"),
        "用户 ID": data.get("userId", "未提供"),
        "会话 ID": data.get("sessionId", "未提供"),
        "复盘版本": data.get("reviewVersion", "未提供"),
        "来源类型": data.get("sourceType", "未提供"),
        "来源会话事件": data.get("sourceSessionEventId", "未提供"),
        "完成时间": data.get("completedAt", "未提供"),
    }
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in identity.items():
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    document.add_heading("二、证据质量", level=1)
    _value(document, {
        "证据类型": data.get("evidenceType", "未提供"),
        "证据置信度": data.get("evidenceConfidence", "未提供"),
        "画像变化是否应用": data.get("applyProfileChanges", False),
        "持久化状态": data.get("persistenceStatus", "未提供"),
    })
    document.add_heading("三、综合评价", level=1)
    overall = data.get("overallAssessment", data.get("overall_assessment", ""))
    _value(document, overall)
    document.add_heading("四、逐题复盘", level=1)
    questions = data.get("questionReviews", [])
    if not isinstance(questions, list) or not questions:
        document.add_paragraph("本场未记录逐题复盘。")
    else:
        for index, item in enumerate(questions, start=1):
            if not isinstance(item, dict):
                continue
            question = item.get("question", item.get("originalQuestion", "未记录问题"))
            document.add_heading(f"{index}. {question}", level=2)
            for label, keys in (
                ("问题编号", ("questionId",)),
                ("回答摘要", ("answer", "answerSummary", "originalAnswer")),
                ("正确性", ("correctness", "technicalCorrectness")),
                ("完整性", ("completeness",)),
                ("错误与遗漏", ("errors", "omissions", "errorsAndOmissions")),
                ("评价", ("evaluation", "assessment")),
                ("参考答案", ("standardAnswer", "referenceAnswer")),
                ("推荐口述", ("spokenAnswer",)),
                ("表达分析", ("expressionAnalysis",)),
                ("变式复测", ("retest", "nextFollowUps")),
            ):
                value = next((item[key] for key in keys if key in item), "未提供")
                paragraph = document.add_paragraph()
                label_run = paragraph.add_run(f"{label}：")
                _font(label_run, bold=True)
                paragraph.add_run(str(value))
    document.add_heading("五、结构化画像变化", level=1)
    changes = data.get("profileChanges", [])
    if isinstance(changes, list):
        for change in changes:
            if isinstance(change, dict):
                _value(document, change)
            else:
                document.add_paragraph(str(change), style="List Bullet")
    else:
        _value(document, changes)
    document.add_heading("六、下一次面试建议", level=1)
    recommendations = data.get("recommendations", [])
    _value(document, recommendations)


def create_review_report(input_path: Path, output_path: Path) -> None:
    """Render the persisted schema-1.2 review JSON and nothing else.

    Legacy report shapes are not rendered here; reading an old report would need a
    separately named read-only adapter, never a branch in the current save path.
    """
    data = json.loads(input_path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("review record JSON root must be an object")
    if data.get("eventType") != "interview.review.completed" or data.get("schemaVersion") != "1.2":
        raise ValueError("review record must be a schemaVersion 1.2 interview.review.completed JSON")

    document = Document()
    _configure(document)
    _schema12_report(document, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    arguments = parser.parse_args()
    create_review_report(arguments.input, arguments.output)


if __name__ == "__main__":
    main()
